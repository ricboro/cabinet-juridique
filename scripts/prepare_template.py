"""
Script one-shot : convertit le template DOCX source (placeholders [X])
en template docxtpl ({{ variable }}) et l'enregistre dans Template/.

Usage : python3 scripts/prepare_template.py
"""
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = Path("Template/Convention d_honoraires generique.docx")
DST = Path("Template/convention_honoraires_template.docx")

# Remplacement au niveau de chaque élément <w:t> individuel
# (gère aussi les hyperlinks et les placeholders dans des runs isolés)
ELEMENT_REPLACEMENTS = {
    "[MAIL]": "{{ mail }}",
    "[TELEPHONE]": "{{ telephone }}",
    "[TAUX_HORAIRE]": "{{ taux_horaire }}",
    "[ESTIMATION_TOTAL]": "{{ estimation_total }}",
    "[NB_HEURE_ESTIMATION]": "{{ nb_heure_estimation }}",
    "[DATE_CREATION]": "{{ date_creation }}",
    "[ADRESSE]": "{{ adresse }}",
    "[PRENOM]": "{{ prenom }}",
    "[NOM]": "{{ nom }}",
    "[A REDIGER]": "[A REDIGER]",  # conservé tel quel
}

# Remplacement au niveau du texte complet du paragraphe
# (pour les placeholders qui couvrent tout le paragraphe)
PARAGRAPH_REPLACEMENTS = {
    "Docteur [PRENOM] [NOM]": "{{ client_line1 }}",
    "Médecin anesthésiste-réanimateur": "{{ client_line2 }}",
}

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def replace_in_paragraph(para) -> bool:
    """
    Remplace les placeholders dans un paragraphe.
    Stratégie 1 : remplacement plein paragraphe (fusionne les runs si besoin).
    Stratégie 2 : remplacement élément par élément <w:t> (hyperlinks, runs isolés).
    """
    modified = False

    # --- Stratégie 1 : remplacement plein paragraphe ---
    full_text = para.text
    new_text = full_text
    for old, new in PARAGRAPH_REPLACEMENTS.items():
        new_text = new_text.replace(old, new)
    if new_text != full_text:
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        modified = True
        return modified

    # --- Stratégie 2 : remplacement par élément <w:t> ---
    for t_elem in para._p.findall(f".//{{{NS}}}t"):
        old = t_elem.text or ""
        new = old
        for src, dst in ELEMENT_REPLACEMENTS.items():
            new = new.replace(src, dst)
        if new != old:
            t_elem.text = new
            modified = True

    return modified


def insert_paragraph_after(ref_para, text) -> None:
    """Insère un nouveau paragraphe après ref_para avec le texte donné."""
    new_p = OxmlElement("w:p")
    if ref_para._p.pPr is not None:
        new_p.append(deepcopy(ref_para._p.pPr))
    new_r = OxmlElement("w:r")
    if ref_para.runs and ref_para.runs[0]._r.rPr is not None:
        new_r.append(deepcopy(ref_para.runs[0]._r.rPr))
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)
    ref_para._p.addnext(new_p)


def main():
    doc = Document(SRC)
    client_line2_para = None

    for para in doc.paragraphs:
        original = para.text.strip()
        changed = replace_in_paragraph(para)
        if changed:
            print(f"  [OK] '{original[:70]}' -> '{para.text.strip()[:70]}'")
        if "{{ client_line2 }}" in para.text:
            client_line2_para = para

    if client_line2_para is not None:
        insert_paragraph_after(client_line2_para, "{{ client_line3 }}")
        print("  [OK] Paragraphe {{ client_line3 }} inséré après client_line2")

    doc.save(DST)

    # Vérification finale
    print("\n--- Vérification placeholders restants ---")
    doc2 = Document(DST)
    found = False
    for i, para in enumerate(doc2.paragraphs):
        t = para.text.strip()
        if "[" in t and "[A REDIGER]" not in t:
            print(f"  [WARN] Para {i}: {t[:80]}")
            found = True
    if not found:
        print("  Aucun placeholder [X] résiduel.")

    print(f"\nTemplate enregistré : {DST}")


if __name__ == "__main__":
    main()
